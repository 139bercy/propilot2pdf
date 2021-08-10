import os
import json
import datetime

import re
import pandas as pd
import numpy as np
from unidecode import unidecode
import docx

# Permet la génération de word
from docx import Document
from docx.shared import Pt
from docxcompose.composer import Composer
from docxtpl import DocxTemplate, RichText
from docx.enum.style import WD_STYLE_TYPE

# Logger
import logging
import logging.handlers
# Définition du logger
logger = logging.getLogger("main.build_reports")
logger.setLevel(logging.DEBUG)


# Variable globale
mailles = ["national", "regional", "departemental"]
volet2code_mesures = {
    'Ecologie': ["MPR2", "MPR4", "BOE1", "DVP1", "RBC3", "RBE1", "AEA1", "FFR1", "BPI1", "BPI2"],  #MPR et BPI x2
    'Compétitivité': ["IDF1", "IDF2", "IDF3", "PIT3", "SAC3", "FUM1", "SFC1", "SBF1"],
    'Cohésion': ["APP1", "PEJ1", "CIE1", "PEC1", "CDP1", "GJE1", "SCI1", "PTH1", "SIL1"],
}
dico_mois = {"Janvier":"01-31",
             "Février":"02-28",
             "Mars":"03-31",
             "Avril":"04-30",
             "Mai":"05-31",
             "Juin":"06-30",
             "Juillet":"07-31",
             "Août":"08-31",
             "Septembre":"09-30",
             "Octobre":"10-31",
             "Novembre":"11-30",
             "Décembre":"12-31"}

months = ('Janvier', 'Février', 'Mars', 'Avril', 'Mai', 'Juin', 'Juillet', 
            'Août', 'Septembre', 'Octobre', 'Novembre', 'Décembre')

# Liste des régions pour lesquelles on ne veut pas de fiche. Noms provenant de taxo_regs.csv
L_reg_no_output = ["00"]  # 00 correspond à Etranger
L_dep_no_output = ['00']
# Dossier imgs avec les logos
img_dir_path = './img/'
word_dir_path = "reports_word"
word_gen_dir_path = "reports_word/Generation_p2p"

# main du fichier
def main_build_reports():
    # Import ref départements / régions
    taxo_dep_df = pd.read_csv('refs/taxo_deps.csv', dtype={'dep':str, 'reg':str})
    taxo_dep_df['dep'] = taxo_dep_df['dep'].apply(lambda x: x.zfill(2))
    taxo_dep_df['reg'] = taxo_dep_df['reg'].apply(lambda x: x.zfill(2))
    dep_list = list(taxo_dep_df['dep'].unique())
    logger.info('{} departements.'.format(len(dep_list)))

    taxo_reg_df = pd.read_csv('refs/taxo_regions.csv', dtype={'reg':str})
    taxo_reg_df['reg'] = taxo_reg_df['reg'].apply(lambda x: x.zfill(2))
    reg_list = list(taxo_reg_df['reg'].unique())
    logger.info('{} regions.'.format(len(reg_list)))

    pp_dep = pd.read_csv("pp_dep.csv", sep=";", dtype={"reg":str, "dep":str})

    # Suppression des espacements multiples dans la date
    pp_dep['Date'] = pp_dep.Date.apply(lambda x: re.sub(' +', ' ', x))
    pp_dep['code_mesure'] = pp_dep.indicateur.apply(lambda x: x.split('-')[-1].strip())

    # Avoir le nom des mesures utilisé dans pp_dep séparés par volet
    volet2mesures = create_dictionnaire_volet2mesures(pp_dep)
     
    #Obtention des valeurs cumulées Régionales et Nationale
    pp_reg, pp_nat = add_cumulated_value(pp_dep, taxo_reg_df)

    # On veut relier mesure -> indicateurs
    dict_mesure_indic = create_dict_mesure_indic(pp_dep, volet2code_mesures)

    # On ne veut pas afficher les lignes de Prime Rénov nulles
    pp_dep = pp_dep.loc[(pp_dep.short_mesure != "Ma Prime Rénov'") | (pp_dep.valeur != 0) ]
    pp_reg = pp_reg.loc[(pp_reg.short_mesure != "Ma Prime Rénov'") | (pp_reg.valeur != 0) ]
    pp_nat = pp_nat.loc[(pp_nat.short_mesure != "Ma Prime Rénov'") | (pp_nat.valeur != 0) ]

    assert pp_dep[(pp_dep['valeur'] == 0) & (pp_dep.short_mesure == "Ma Prime Rénov'")].shape[0] == 0
    assert pp_reg[(pp_reg['valeur'] == 0) & (pp_reg.short_mesure == "Ma Prime Rénov'")].shape[0] == 0
    assert pp_nat[(pp_nat['valeur'] == 0) & (pp_nat.short_mesure == "Ma Prime Rénov'")].shape[0] == 0

    assert pp_nat.duplicated(subset=['mesure','short_indic', 'Date']).sum() == 0
    assert pp_reg.duplicated(subset=['mesure','short_indic', 'Date', 'reg']).sum() == 0

    # Calcul des poids dep/reg
    pp_dep, pp_reg = add_weighted_value(pp_dep, pp_reg, pp_nat)

    # formatage de la variable valeur
    pp_dep.valeur = pp_dep.valeur.astype(str)
    pp_dep.valeur = pp_dep.apply(lambda x: str(format_amount(x["short_indic"], x["valeur"])) + ' (' + x['poids_reg'] + ')', axis=1)
    pp_reg.valeur = pp_reg.valeur.astype(str)
    pp_reg.valeur = pp_reg.apply(lambda x: str(format_amount(x["short_indic"], x["valeur"])) + ' (' + x['poids_nat'] + ')', axis=1)
    pp_nat.valeur = pp_nat.valeur.astype(str)
    pp_nat.valeur = pp_nat.apply(lambda x: format_amount(x["short_indic"], x["valeur"]), axis=1)

    # Création des dossiers pour stocker les fiches
    
    # Dossiers fiches
    mkdir_ifnotexist(word_dir_path)
    mkdir_ifnotexist(word_gen_dir_path)
    all_charts_as_df = {"departemental": {dep: {} for dep in dep_list},
                    "national": {'France': {}},
                    "regional": {reg: {} for reg in reg_list}}
    # Récuperer les 3 derniers mois à insérer dans les fiches
    last_dates_to_keep = insert_months_to(modulo=1)

    make_all_charts(dict_mesure_indic, pp_dep, pp_reg, pp_nat, taxo_dep_df, last_dates_to_keep, months, all_charts_as_df)
    check_charts_exhaustivity(all_charts_as_df, taxo_dep_df, taxo_reg_df, dict_mesure_indic)
    short_mesure2url, short_mesure2to_comment = insert_mesure_to(pp_dep)
    
    # Lance la génération dans le dossier reports_word
    create_all_dep(taxo_dep_df, taxo_reg_df, volet2mesures, all_charts_as_df, short_mesure2to_comment, short_mesure2url, dict_mesure_indic)
    check_num_docx_created(taxo_dep_df)    

# Fonction nécessaire

def create_dict_mesure_indic(pp_dep: pd.DataFrame, volet2code_mesures: dict) -> dict:
    """
    Creates crossing dictionnary between Mesure (key) and Indicateur (value)
    """
    # Extraction des mesures-indicateurs à afficher dans les fiches
    code_mesures_to_keep = set([mesure for volet in volet2code_mesures for mesure in volet2code_mesures[volet]])
    mesure_indics = pp_dep.groupby(['code_mesure', 'short_mesure']).agg({'short_indic': list}).reset_index()
    mesure_indics = mesure_indics[mesure_indics.code_mesure.isin(code_mesures_to_keep)]
    dict_mesure_indic = {}

    for i, row in mesure_indics.iterrows():
        dict_mesure_indic[row['short_mesure']] = list(set(row['short_indic']))

    # On se restreint à certains indicateurs dans les mesures suivantes
    dict_mesure_indic['Soutien aux fonds propres des filières automobiles et aéronautiques'] = ["Nombre d'entreprises"]
    dict_mesure_indic['AAP Industrie : Soutien aux projets industriels territoires'] = ['Nombre de TPE,PME,ETI bénéficiaires']
    dict_mesure_indic['AAP Industrie : Sécurisation approvisionnements critiques'] = ['Nombre de TPE,PME,ETI bénéficiaires']

    # Rajout de restriction ICI
    dict_mesure_indic["MaPrimeRénov'"] = ['Nombre de dossiers MaPrimeRénov validés', 'Montant total des travaux associés aux dossiers validés']
    return dict_mesure_indic


def create_dictionnaire_volet2mesures(pp_dep: pd.DataFrame) -> dict:
    """
    Create crossing dictionnary between Mesure (key) and Indicateur (value)
    """
    volet2mesures = {volet: [] for volet in volet2code_mesures}
    for volet in volet2code_mesures:
        # Trier les mesures par ordre alphabétique
        mesures = pp_dep[pp_dep.code_mesure.isin(volet2code_mesures[volet])].short_mesure.sort_values().unique().tolist()
        volet2mesures[volet] = mesures
    return volet2mesures


def add_cumulated_value(pp_dep: pd.DataFrame, taxo_reg_df: pd.DataFrame) -> list: 
    """
    Create 2 new dataframe: pp_reg and pp_nat with cumulated value into valeur
    
    Returns:
        list[0]: pp_reg
        list[1]: pp_nat
    """
    # Obtention des valeurs régionale par somme des valeurs départementales
    pp_reg = pd.pivot_table(pp_dep, index=["mesure","short_mesure", "reg","region", "Date", "period_date", "short_indic"], values="valeur", aggfunc=np.sum)
    pp_reg.rename(columns={"reg":"libelle"}, inplace=True)
    pp_reg.reset_index(inplace=True)
    check_pp_reg(pp_reg, taxo_reg_df, pp_dep)

    # Obtention des valeurs nationales par somme des valeurs régionale
    pp_nat = pd.pivot_table(pp_reg, index=["mesure", "short_mesure", "Date","period_date", "short_indic"], values="valeur", aggfunc=np.sum)
    pp_nat.reset_index(inplace=True)
    check_pp_nat(pp_nat)
    return pp_reg, pp_nat


def insert_months_to(modulo: int = 0, months: tuple = months, nb_mois: int = 3) -> list:
    """
    Collect the 3 months to insert in the parlementary files
    """
    # modulo: Mois en cours - Modulo = dernier mois présent sur la fiche
    # months: couple de mois, variable globale définie en début de script
    # nb_mois: Combien de mois apparaitront sur la fiche 3 par défaut
    today = datetime.date.today()
    last_dates_to_keep = []
    for i in range(1, nb_mois+1):
        month_name = months[(today.month-modulo-i) % 12]
        year = today.year - 1 if (today.month-1-i) < 0 else today.year
        last_dates_to_keep.append(f'{month_name} {year}')
    return last_dates_to_keep


def add_weighted_value(pp_dep: pd.DataFrame, pp_reg: pd.DataFrame, pp_nat: pd.DataFrame) -> list:
    """
    Add weigted_value dep/reg into pp_dep and reg/nat into pp_reg 
    
    Returns:
        list[0]: pp_dep
        list[1]: pp_reg
    """
    # Calcul des poids dep/reg
    pp_dep = pp_dep.merge(pp_reg[['mesure','short_indic', 'Date', 'reg', 'valeur']], 
                        on=['mesure','short_indic', 'Date', 'reg'], 
                        how='left', suffixes=('', '_reg'))
    pp_dep['poids_reg'] = pp_dep.apply(lambda x: str(round(100 * x['valeur'] / max(x['valeur_reg'], 1))) + "%", axis=1)

    # Vérifier qu'on a pas de pourcentages aberrants
    assert pp_dep.poids_reg.isnull().sum() == 0
    assert all(int(poids_reg[:-1]) <= 100 for poids_reg in pp_dep.poids_reg.values)

    # Calcul des poids reg/nat
    pp_reg = pp_reg.merge(pp_nat[['mesure','short_indic', 'Date', 'valeur']],
                        on=['mesure','short_indic', 'Date'], 
                        how='left', suffixes=('', '_nat'))
    pp_reg['poids_nat'] = pp_reg.apply(lambda x: str(round(100 * x['valeur'] / max(1, x['valeur_nat']))) + "%", axis=1)

    # Vérifier qu'on a pas de pourcentages aberrants
    assert pp_reg.poids_nat.isnull().sum() == 0
    assert all(int(poids_nat[:-1]) <= 100 for poids_nat in pp_reg.poids_nat.values)
    return pp_dep, pp_reg


def mkdir_ifnotexist(path: str):
    """
    Create a folder if it's doesn't exist
    """
    if not os.path.isdir(path):
        os.mkdir(path)


def format_thousands(s: str) -> str: 
    """
    Format a number like 1000 into 1 000
    """
    new_str = ''
    for i, ch in enumerate(s[::-1], start=1):
        new_str = ch + new_str
        if i % 3 == 0:
            new_str = ' ' + new_str
    return new_str.strip()


def format_amount(indic: str, valeur: str) -> str:
    """
    Add unit after amount
    """
    if "Montant" in indic:
        f_valeur = float(valeur)
        if f_valeur > 1000000:
            return str(round(f_valeur/1000000, 1)) + ' M€'
        elif f_valeur > 10000:
            return str(round(f_valeur/1000, 1)) + ' k€'
        else:
            return str(f_valeur)
    else:
        try:
            return format_thousands(valeur.split(".")[0])
        except ValueError as err:
            logger.info(f"L'indicateur {indic} possède des valeurs invalides : {err}")


def check_pp_reg(pp_reg: pd.DataFrame, taxo_reg_df: pd.DataFrame, pp_dep: pd.DataFrame):
    """
    Check if we have all regions and measures in pp_reg
    """
    assert sorted(pp_reg['reg'].unique()) == sorted(taxo_reg_df['reg'])
    assert sorted(pp_reg['region'].unique()) == sorted(taxo_reg_df['libelle'])
    assert sorted(pp_reg['mesure'].unique()) == sorted(pp_dep['mesure'].unique())
    assert sorted(pp_reg['short_mesure'].unique()) == sorted(pp_dep['short_mesure'].unique())
    

def check_pp_nat(pp_nat: pd.DataFrame):
    """
    Check if we have all measures in pp_nat
    """
    assert sorted(pp_nat['mesure'].unique()) == sorted(pp_nat['mesure'].unique())
    assert sorted(pp_nat['short_mesure'].unique()) == sorted(pp_nat['short_mesure'].unique())


def complete_values_for_missing_dates(df_plot: pd.DataFrame, na_replacement: str, last_dates_to_keep: list, months: tuple) ->pd.DataFrame:
    """
    For each date in last_dates_to_keep, if the date is missing in df_plot then we fill with na_replacement
    """
    # Obtention des dates à ajouter
    missing_dates = set(last_dates_to_keep) - set(df_plot.Date)
    sorted_missing_dates = sorted(missing_dates, key=lambda x: months.index(x.split(' ')[0]))

    # Récupération des 3 derniers mois
    dict_conv = {}
    for i in range(len(last_dates_to_keep)-1, -1, -1):
        mois, annee = last_dates_to_keep[i].split(" ")
        dict_conv[last_dates_to_keep[i]] = str(annee) + "-" + dico_mois[mois] + 'T00:00:00.0000000'

    # Remplissage pour les dates manquantes
    df_complement = pd.DataFrame({col: sorted_missing_dates if col == 'Date' else na_replacement for col in df_plot.columns})

    df_complement['period_date'] = '2021-05-31T00:00:00.0000000'
    if len(sorted_missing_dates) != 0:
        for element in sorted_missing_dates:
            df_complement['period_date'] = np.where(df_complement.Date == element, dict_conv[element], df_complement['period_date'])

    return pd.concat([df_plot, df_complement]).reset_index(drop=True)


def make_pp_chart(maille: str, mesure: str, short_indics: list, pp_dep: pd.DataFrame, pp_reg: pd.DataFrame,
                  pp_nat: pd.DataFrame, taxo_dep_df: pd.DataFrame, last_dates_to_keep: list, months: tuple, all_charts_as_df: dict):
    """
    For one mesure and one maille, create every table we need to generate parlementary file, and stock them in all_charts_as_df.  
    Depending on the parameter maille, tables will be create:
        - with pp_dep if maille = départemental, for all department
        - with pp_reg if maille = régional, for all region
        - pp_nat otherwise
    Below, layout of one created table with pp_dep: 
     
            -----------------------
            |        maille:       |
            ------------------------
            |shrtind1|....|shrtindn|
            ------------------------
       date1|   data from pp_dep   |
       date2|   data from pp_dep   |
       date3|   data from pp_dep   |
            ------------------------
    """
    na_replacement = "Indisponible"
    
    if maille == "departemental":
        df = pp_dep.loc[(pp_dep.short_mesure == mesure)].sort_values(by="period_date", ascending=True).copy()
        deps = taxo_dep_df.dep.unique()  # Liste exhaustive de départements
        
        # Préparer un tableau par défaut à mettre quand on ne dispose d'aucune valeur
        default = df.groupby(["Date", "period_date"]).sum().sort_values("period_date", ascending=True).reset_index()
        default[short_indics] = na_replacement
        default = default[["Date", "period_date"] + short_indics]
        default = complete_values_for_missing_dates(default, na_replacement, last_dates_to_keep, months)
        default = default.reset_index()
        default = default.sort_values(by = 'period_date', ascending=True)
        default = default.drop('period_date', axis=1)
        default = default.drop('index', axis=1)

        for dep in deps:
            df_dep = df.loc[df.dep == dep]
            if df_dep.shape[0] == 0:
                all_charts_as_df[maille][dep][mesure] = default.T.reset_index().T  # Avoir le nom des colonnes en valeurs
            else:
                df_plot = pd.pivot_table(df_dep, index=['period_date', 'Date'], columns=['short_indic'], values='valeur', aggfunc='first')
                df_plot = df_plot.reset_index().sort_values(by = 'period_date')
                df_plot = df_plot.rename_axis(None, axis=1)
                df_plot = df_plot.fillna(na_replacement)
                # Ajout des indicateurs/colonnes manquantes
                cols = set(df_plot.columns).intersection(short_indics)
                if len(cols) != len(short_indics):
                    missing_cols = set(short_indics) - cols
                    for missing_col in missing_cols:
                        df_plot[missing_col] = na_replacement
                df_plot = df_plot[['Date', 'period_date'] + short_indics]
                df_plot = complete_values_for_missing_dates(df_plot, na_replacement, last_dates_to_keep, months)
                df_plot = df_plot.reset_index().sort_values(by = 'period_date')
                df_plot = df_plot.drop('period_date', axis=1)
                df_plot = df_plot.drop('index', axis=1)
                all_charts_as_df[maille][dep][mesure] = df_plot.T.reset_index().T
                
            
    elif maille == "regional":
        df = pp_reg.loc[(pp_reg.short_mesure == mesure)].sort_values(by="period_date", ascending=True).copy()
        regs = taxo_dep_df.reg.unique()
        
        default = df.groupby(["Date", "period_date"]).sum().sort_values("period_date", ascending=True).reset_index()
        default[short_indics] = na_replacement
        default = default[["Date", "period_date"] + short_indics]
        default = complete_values_for_missing_dates(default, na_replacement, last_dates_to_keep, months)
        default = default.reset_index()
        default = default.sort_values(by = 'period_date', ascending=True)
        default = default.drop('period_date', axis=1)
        default = default.drop('index', axis=1)

        for reg in regs:
            df_reg = df.loc[df.reg == reg]
            if df_reg.shape[0] == 0:
                all_charts_as_df[maille][reg][mesure] = default.T.reset_index().T
            else:
                df_plot = pd.pivot_table(df_reg, index=['period_date', 'Date'], columns=['short_indic'], values='valeur', aggfunc='first')
                df_plot = df_plot.reset_index()
                df_plot = df_plot.rename_axis(None, axis=1)
                df_plot = df_plot.fillna(na_replacement)
                cols = set(df_plot.columns).intersection(short_indics)
                if len(cols) != len(short_indics):
                    missing_cols = set(short_indics) - cols
                    for missing_col in missing_cols:
                        df_plot[missing_col] = na_replacement
                df_plot = df_plot[['Date', 'period_date'] + short_indics]
                df_plot = complete_values_for_missing_dates(df_plot, na_replacement, last_dates_to_keep, months)
                df_plot = df_plot.reset_index().sort_values(by = 'period_date')
                df_plot = df_plot.drop('period_date', axis=1)
                df_plot = df_plot.drop('index', axis=1)
                all_charts_as_df[maille][reg][mesure] = df_plot.T.reset_index().T
            
    elif maille == "national":
        df_nat = pp_nat.loc[(pp_nat.short_mesure == mesure)].sort_values(by="period_date", ascending=True).copy()
        df_plot = pd.pivot_table(df_nat, index=['period_date', 'Date'], columns=['short_indic'], values='valeur', aggfunc='first')
        df_plot = df_plot.reset_index()
        df_plot = df_plot.rename_axis(None, axis=1)
        df_plot = df_plot.fillna(na_replacement)
        df_plot = df_plot[['Date', 'period_date'] + short_indics]
        df_plot = complete_values_for_missing_dates(df_plot, na_replacement, last_dates_to_keep, months)
        df_plot = df_plot.reset_index().sort_values(by = 'period_date')
        df_plot = df_plot.drop('period_date', axis=1)
        df_plot = df_plot.drop('index', axis=1)
        all_charts_as_df[maille]['France'][mesure] = df_plot.T.reset_index().T


def make_all_charts(dict_mesure_indic: dict, pp_dep: pd.DataFrame, pp_reg: pd.DataFrame, pp_nat: pd.DataFrame, taxo_dep_df: pd.DataFrame,
                    last_dates_to_keep: list, months: tuple, all_charts_as_df: dict):
    """
    Create all tables we need to generate every parlementary file and store them in all_charts_as_df
    """
    for mesure in dict_mesure_indic:
        short_indics = dict_mesure_indic[mesure]
        for maille in mailles :
            make_pp_chart(maille, mesure, short_indics, pp_dep, pp_reg, pp_nat, taxo_dep_df, last_dates_to_keep, months, all_charts_as_df)


def check_charts_exhaustivity(all_charts_as_df: dict, taxo_dep_df: pd.DataFrame, taxo_reg_df: pd.DataFrame, dict_mesure_indic: dict):
    """
    Multiple checks
            - Check if all_charts_as_df contains all departments, regions and the key France
            - Check if all_charts_as_df contains all tables for each department, region and for the key France
    """
    assert sorted(all_charts_as_df['departemental'].keys()) == sorted(taxo_dep_df['dep'])
    assert sorted(all_charts_as_df['regional'].keys()) == sorted(taxo_reg_df['reg'])
    assert sorted(all_charts_as_df['national'].keys()) == ['France']
    
    # Vérifier si des graphiques manquent.
    for dep in taxo_dep_df['dep']:
        assert sorted(all_charts_as_df['departemental'][dep].keys()) == sorted(dict_mesure_indic.keys()), f"{dep}"
    for reg in taxo_reg_df['reg']:
        assert sorted(all_charts_as_df['regional'][reg].keys()) == sorted(dict_mesure_indic.keys())
    
    assert sorted(all_charts_as_df['national']['France'].keys()) == sorted(dict_mesure_indic.keys())

def insert_mesure_to(pp_dep: pd.DataFrame) -> tuple:
    """ 
    Only some measures will be present in parlementary file and for the latter 
    we will keep an URL for extra information, and if we have to create a commentary space in the file 
    
    Returns:
        list[0]: short_mesure2url
        list[1]: short_mesure2to_comment
    """
    # Importer le dataframe des mesures à insérer
    ref_mesures2 = pd.read_excel('refs/20210630_Liste_Mesures-Ficheparlementaire.xlsx')
    ref_mesures2.drop(["Unnamed: 5", "Mesures suivie dans le TdB grand public"], axis = 1, inplace=True)
    ref_mesures2.drop([27], inplace=True)
    ref_mesures2 = ref_mesures2.rename(columns={"Liens hypertexte": "url",
                                                "Numéro indicateur": "code_mesure"})
    for i in range(27):
        if i < 10:
            ref_mesures2["Volet"].loc[i] = ref_mesures2["Volet"].loc[0]
        elif i < 18:
            ref_mesures2["Volet"].loc[i] = ref_mesures2["Volet"].loc[10]
        else:
            ref_mesures2["Volet"].loc[i] = ref_mesures2["Volet"].loc[18]

    ref_mesures2["Mesures"].iloc[1] = ref_mesures2["Mesures"].iloc[0]
    ref_mesures2["Mesures"].iloc[9] = ref_mesures2["Mesures"].iloc[8]
    ref_mesures2["Mesures"].iloc[11] = ref_mesures2["Mesures"].iloc[10]
    ref_mesures2["Mesures"].iloc[12] = ref_mesures2["Mesures"].iloc[10]

    ref_mesures2["url"].iloc[1] = ref_mesures2["url"].iloc[0]
    ref_mesures2["url"].iloc[9] = ref_mesures2["url"].iloc[8]
    ref_mesures2["url"].iloc[11] = ref_mesures2["url"].iloc[10]
    ref_mesures2["url"].iloc[12] = ref_mesures2["url"].iloc[10]

    ref_mesures2["url"].iloc[16] = "https://www.economie.gouv.fr/files/files/directions_services/plan-de-relance/Guide-mesures-relance-exportations.pdf"
    ref_mesures2["url"].iloc[17] = "https://www.economie.gouv.fr/files/files/directions_services/plan-de-relance/Guide-mesures-relance-exportations.pdf"

    ref_mesures2["code_mesure"].iloc[18] = "APP"
    L_com = ["RBE",
            "FAA",
            "PIT",
            "SAC",
            "FUM",
            "SBF"]
    ref_mesures2["commentaire"] = "n"
    for i in range(27):
        for j in range(len(L_com)):
            if L_com[j] in ref_mesures2["code_mesure"].iloc[i]:
                ref_mesures2["commentaire"].iloc[i] = "o"

    # Retravail de la colonne code_mesure
    code = ["MPR4", "MPR2", "BOE1", "DVP1", "RBC3", "RBE1", "AEA1", "FFR1", "BPI1", "BPI2", "IDF3", "IDF1", "IDF2", "PIT3", "SAC3", "FUM1", "SFC1", "SBF1",
    "APP1", "PEJ1", "CIE1", "PEC1", "CDP1", "GJE1", "SCI1", "PTH1", "SIL1"]

    for i in range(len(code)):
        ref_mesures2["code_mesure"].iloc[i] = code[i]
    ref_mesures = ref_mesures2

    code2short_mesure = pp_dep[['code_mesure', 'short_mesure']].set_index('code_mesure').to_dict()['short_mesure']
    # Ajout des clés manquantes
    code2short_mesure['SIL1'] = "Soutien à l'investissement local (DSIL exceptionnelle)"
    code2short_mesure['RBC3'] = "Rénovation thermique des bâtiments publics soutenus par la DSIL, DSID, DRI"

    ref_mesures['short_mesure_in_pp_dep'] = ref_mesures.code_mesure.apply(lambda x: code2short_mesure[x])
    short_mesure2url = ref_mesures.groupby('short_mesure_in_pp_dep').agg({'url': list}).apply(lambda x: x['url'][0].strip(), axis=1).to_dict()
    short_mesure2to_comment = ref_mesures.groupby('short_mesure_in_pp_dep').agg({'commentaire': list}).apply(lambda x: x['commentaire'][0] == 'o', axis=1).to_dict()

    return short_mesure2url, short_mesure2to_comment


def create_front_page(nom_departement: str) -> str:
    """
    Given a department, creates the front page of the parlementary file
    """
    doc = DocxTemplate("template/template_front_page.docx")
    today = datetime.datetime.today()
    months = ('Janvier', 'Février', 'Mars', 'Avril', 'Mai', 'Juin', 'Juillet', 
            'Août', 'Septembre', 'Octobre', 'Novembre', 'Décembre')
    today_str = f"{months[today.month-1]} {today.year}"
    context = {'dep': str(nom_departement), 
               'date': today_str}  # A remplacer par today_str plus tard. On nous demande de mettre Mai 2021 ------------------------------------ !!!!!!!!!!!!!!!!!!!!
    doc.render(context)
    name_file = "reports_word/Generation_p2p/front_page_{}.docx".format(nom_departement)
    doc.save(name_file)
    return name_file


def normalize_name(name: str) -> str:
    """
    Normalize a str: delete whitespace, special characters, put in lowercase and keep only letters
    """
    name = name.lower()
    name = unidecode(name)
    name = re.sub('[^a-z]', ' ',  name)
    name = re.sub(' +', '', name)
    return name


def create_volet_page(nom_volet: str, num_volet: int) -> str:
    """
    Create for one departement, the "volet" page of the parlementary file
    """
    doc = DocxTemplate("template/template_volet.docx")
    context = {'volet': nom_volet, 
               'num_volet': num_volet, 
               'code_comment': "{% for f in " + normalize_name(nom_volet) + " %}{{ f.text }} {{ f.image }} {% endfor %}",
              }
    doc.render(context)
    name_file = "reports_word/Generation_p2p/{}.docx".format(nom_volet)
    doc.save(name_file)
    return name_file


def delete_paragraph(paragraph: docx.text.paragraph.Paragraph):
    """
    Delete the paragraph
    """
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None
    
    
def fusion_word(word1: str, word2: str, dep: str) -> str:
    """
    Concatenate word2 at the end of word1
    """
    master = Document(word1)
    master.add_page_break()
    composer = Composer(master)
    doc1 = Document(word2)
    composer.append(doc1)
    name_fusion = "reports_word/Suivi_Territorial_plan_relance_{}.docx".format(dep)
    composer.save(name_fusion)
    return name_fusion


def create_content_page(all_charts_as_df: dict, departement: str, region: str, mesure: str, volet: str, dep_name: str, reg_name: str,
                        num_mesure: int, short_mesure2to_comment: dict, short_mesure2url: dict, dict_mesure_indic: dict) -> str:
    """
    Create one page for parlementary file. This page is made of 3 tables (departmental, regional, national) for one measure 
    """
    # Ouverture de template
    if mesure in short_mesure2to_comment and short_mesure2to_comment[mesure]:
        doc = DocxTemplate("template/template_content_page.docx")
    else:
        doc = DocxTemplate("template/template_content_page_no_comment.docx")
    # Recuperation des datas pour les 3 scales
    df_nat = all_charts_as_df["national"]["France"][mesure]
    df_reg = all_charts_as_df["regional"][region][mesure]
    df_dep = all_charts_as_df["departemental"][departement][mesure]
    
    # Recuperation des noms des colonnes
    col_labels = df_nat.iloc[0]
    short_indic = dict_mesure_indic[mesure][0]
    rt_hyperlien = RichText(f"{num_mesure} - ", font='Marianne', size=40, color='#00a65d')
    
    # Si pas d'url trouvé pour la mesure, on redirige le lecteur vers la page de recherche decommenter lors de la maj du xlsx
    url = short_mesure2url[mesure]
    rt_hyperlien.add(f'{mesure}', url_id=doc.build_url_id(url), 
                                  underline=True, color='#00a65d',
                                  font='Marianne', size=40)
    context = {
                'mesure': rt_hyperlien,        
                'title_table_nat' : "Niveau National", 
                'title_table_reg' : "Niveau Régional", 
                'title_table_dep' : "Niveau Départemental",
                'lib_reg': reg_name, 
                'lib_dep': dep_name,
                'col_labels' : col_labels, 
                # Les 3 lignes suivantes permettent de générer des tabeaux avec uniquement les 3 derniers mois.
                # Prend en compte le cas ou il n'y a pas encore 3 mois de données
                'tbl_contents_nat': [{'cols' : list(df_nat.iloc[-i-1])} for i in range(min(len(df_nat)-1, 3))],
                'tbl_contents_reg': [{'cols' : list(df_reg.iloc[-i-1])} for i in range(min(len(df_reg)-1, 3))],
                'tbl_contents_dep': [{'cols' : list(df_dep.iloc[-(i+1)])} for i in range(min(len(df_dep)-1, 3))],
                'code_comment': "{% for f in " + normalize_name(mesure) + " %}{{ f.text }} {{ f.image }} {% endfor %}",
                }
    doc.render(context)
    name_file = "reports_word/Generation_p2p/content_page_{}.docx".format(mesure)
    doc.save(name_file)
    return name_file


def create_fiche(dep, taxo_dep_df: pd.DataFrame, taxo_reg_df: pd.DataFrame, volet2mesures: dict, all_charts_as_df: dict,
                 short_mesure2to_comment: dict, short_mesure2url: dict, dict_mesure_indic: dict) -> str:
    """
    For one department: Create all pages we need ( front page, "volet" pages and measures pages) and concatenate them into one docx document
    """
    #departement: code departement 01:
    #On a les variables volet2mesures, all_charts
    reg = taxo_dep_df[taxo_dep_df['dep'] == dep].iloc[0]['reg']  # Code region
    if reg in L_reg_no_output:
        return False
    num_volet, num_mesure = 1, 1
    reg_name = taxo_reg_df[taxo_reg_df['reg'] == reg].iloc[0]['libelle']  # libelle
    dep_name = taxo_dep_df[taxo_dep_df['dep'] == dep].iloc[0]['libelle']
    name_fusion = create_front_page(dep_name)
    for volet in list(volet2mesures.keys()):  # 3 itérations, dep_name, reg_name
        name_volet = create_volet_page(volet, num_volet) ##
        name_fusion = fusion_word(name_fusion, name_volet, dep_name) ##
        liste_mesure = volet2mesures[volet]
        num_volet += 1
        for mesure in liste_mesure:
            name_content = create_content_page(all_charts_as_df, dep, reg, mesure, volet, dep_name, reg_name, num_mesure, short_mesure2to_comment, short_mesure2url, dict_mesure_indic) ##
            name_fusion = fusion_word(name_fusion, name_content, dep_name) ##
            num_mesure += 1
    return name_fusion
            


def create_all_dep(taxo_dep_df: pd.DataFrame, taxo_reg_df: pd.DataFrame, volet2mesures: dict, all_charts_as_df: dict,
                   short_mesure2to_comment: dict, short_mesure2url: dict, dict_mesure_indic: dict):
    """
    Create parlementary file for each french department
    """
    list_all_dep = taxo_dep_df[~taxo_dep_df["dep"].isin(L_dep_no_output)].dep
    for dep in list_all_dep:
        docx_path = create_fiche(dep, taxo_dep_df, taxo_reg_df, volet2mesures, all_charts_as_df, short_mesure2to_comment, short_mesure2url, dict_mesure_indic)
        logger.info(dep + ' ' + docx_path)
        doc = Document(docx_path)

        # Retirer le double espacement après les tableaux
        for paragraph in doc.paragraphs:
            # Retirer les lignes vides après les tableaux en fin de page : elles sont reconnues par la taille 
            # de la police <= 2 où ne possède pas de run.
            if paragraph.text.__len__() == 0 and (any(run.font.size <= Pt(2) for run in paragraph.runs if run.font.size is not None) or 
                                                  (all(run.font.size is None for run in paragraph.runs))):
                delete_paragraph(paragraph)
        doc.save(docx_path)
        
        # Réduire la taille du paragraphe après le dernier tableau de chaque page
        # Cela permet d'éviter de créer une nouvelle page quand le texte devient trop long
        doc = Document(docx_path)
        styles = doc.styles
        style = styles.add_style('Custom_style2', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(2)
        for paragraph in doc.paragraphs:
            if len(paragraph.text) == 1:
                paragraph.style = doc.styles['Custom_style2']
        doc.save(docx_path)


def check_num_docx_created(taxo_dep_df: pd.DataFrame):
    """
    Check if we have 109 parlementary files
    """
    # Vérifier si on a bien toutes les fiches
    num_test = len([fn for fn in os.listdir('reports_word') if "Suivi" in fn])
    num_true = taxo_dep_df['dep'].shape[0]
    # num_true-1 car on enlève le département Etranger "00"
    assert num_test == num_true - 1, f"{num_test} -- {num_true}"


if __name__ == "__main__":
    main_build_reports()
