import datetime
from docx import Document
from docxcompose.composer import Composer
from docxtpl import DocxTemplate
import pandas as pd
import time

####### Import depuis le jupyter
import pickle
with open("allcharts.pkl", 'rb') as f:
    all_charts_as_df = pickle.load(f)

pp_dep = pd.read_csv("pp_dep.csv", sep=";", dtype={"reg":str, "dep":str})
taxo_dep_df = pd.read_csv("refs/taxo_deps.csv", dtype={"dep":str})
taxo_reg_df = pd.read_csv("refs/taxo_regions.csv", dtype={"reg":str})
# Liste des régions pour lesquelles on ne veut pas de fiche. Noms provenant de taxo_regs.csv
L_reg_no_output = ["0", "989", "988"]  # 0 correspond à Etranger

volet2mesures = {
'Ecologie': ["Ma Prime Rénov'",
		"Bonus électrique",
		'AAP et AMI Efficacité énergétique',
		'Prime à la conversion des véhicules légers',
		'Soutien recherche aéronautique civil', 
		'Rénovation des bâtiments Etats (marchés notifiés)',],

'Compétitivité': ['Assurance prospection', 
		'France Num : aide à la numérisation des TPE,PME,ETI',
		"AAP Industrie : Soutien aux projets industriels territoires",
		"AAP Industrie : Sécurisation approvisionnements critiques",
		"AAP industrie : Modernisation des filières auto et aéro",
		"Renforcement subventions Business France",],

'Cohésion': ["Apprentissage",
		"Prime à l'embauche des jeunes",
		"Prime à l'embauche pour les travailleurs handicapés",
		"Contrats Initiatives Emploi (CIE) Jeunes",
		'Contrats de professionnalisation',
		'Garantie jeunes',
		"Parcours emploi compétences (PEC) Jeunes",
		"Service civique",]
}


dict_mesure_indic = {'AAP Industrie : Soutien aux projets industriels territoires': ['Nombre de TPE,PME,ETI bénéficiaires'],
'AAP Industrie : Sécurisation approvisionnements critiques': ['Nombre de TPE,PME,ETI bénéficiaires'],
'AAP et AMI Efficacité énergétique': ["Nombre d'entreprises ayant reçu l'aide"],
'AAP industrie : Modernisation des filières auto et aéro': ['Nombre de PME'],
'Apprentissage': ['Nombre de contrats d’apprentissage'],
'Assurance prospection': ['Nombre de TPE,PME,ETI bénéficiaires'],
'Bonus électrique': ['Nombre de bonus octroyés à des véhicules électriques'],
'Contrats Initiatives Emploi (CIE) Jeunes': ['Entrées de jeunes en CIE'],
'Contrats de professionnalisation': ['Nombre de contrats de professionnalisation'],
'France Num : aide à la numérisation des TPE,PME,ETI': ["Nombre d'accompagnements dispensés"],
'Garantie jeunes': ['Entrées en garanties jeunes'],
"Ma Prime Rénov'": ['Montant total des primes versées',
'Montant total des travaux',
"Nombre de dossiers MaPrimeRénov' payés"],
'Parcours emploi compétences (PEC) Jeunes': ['Entrées de jeunes en PEC'],
"Prime à l'embauche des jeunes": ["Nombre d'aides à l'embauche des jeunes"],
"Prime à l'embauche pour les travailleurs handicapés": ["Nombre d'aides à l'embauche des travailleurs handicapés"],
'Prime à la conversion des véhicules légers': ['Nombre de primes à la conversion'],
'Renforcement subventions Business France': ["Nombre d'entreprises bénéficiares",
'Nombre de TPE,PME,ETI bénéficiaires'],
'Rénovation des bâtiments Etats (marchés notifiés)': ['Nombre de bâtiments dont le marché de rénovation est notifié'],
'Service civique': ["Nombre d'entrées en service civique"],
'Soutien recherche aéronautique civil': ['Nombre de projets soutenus']}

def get_kpi(dep, short_indic, short_mesure):
    kpi_dep = (pp_dep.loc[(pp_dep.dep == dep) 
                          & (pp_dep.short_mesure == short_mesure) 
                          & (pp_dep.short_indic == short_indic)]
                .sort_values(by="period_date", ascending=False))
    if kpi_dep.shape[0] != 0:
        date= kpi_dep.iloc[0].Date
        valeur = kpi_dep.iloc[0].valeur
    else:
        date = pp_dep.Date.max()
        valeur = 0
    return date, valeur


##### Fin des imports jupyter




def creation_front_page(nom_departement):
    doc = DocxTemplate("template/template_front_page.docx")
    today = datetime.datetime.today().strftime('%Y-%m-%d')
    context = {'dep': str(nom_departement), 
               'date': today}
    doc.render(context)
    name_file = "reports/test/front_page_{}.docx".format(nom_departement)
    doc.save(name_file)
    return name_file


def creation_volet_page(nom_volet):
    doc = DocxTemplate("template/template_volet.docx")
    context = {'volet': str(nom_volet)}
    doc.render(context)
    name_file = "reports/test/{}.docx".format(nom_volet)
    doc.save(name_file)
    return name_file

def fusion_word(word1, word2, dep):
    doc1 = Document(word1)
    doc2 = Document(word2)
    doc1.add_page_break()
    for element in doc2.element.body:
        doc1.element.body.append(element)
    doc1.add_page_break()
    doc1.save("reports_word/Suivi_Territorial_plan_relance_{}.docx".format(dep))
    name_fusion = "reports_word/Suivi_Territorial_plan_relance_{}.docx".format(dep)
    return name_fusion


def fusion2(word1, word2, dep):
    master = Document(word1)
    master.add_page_break()
    composer = Composer(master)
    doc1 = Document(word2)
    composer.append(doc1)
    name_fusion = "reports_word/Suivi_Territorial_plan_relance_{}.docx".format(dep)
    composer.save(name_fusion)
    return name_fusion


def creation_content_page(all_charts_as_df, departement, region, mesure, dep_name, reg_name):
    # Ouverture de template
    doc = DocxTemplate("template/template_content_page_ligne.docx")
    # Recuperation des datas pour les 3 scales
    df_nat = all_charts_as_df["national"]["France"][mesure]
    df_reg = all_charts_as_df["regional"][region][mesure]
    df_dep = all_charts_as_df["departemental"][departement][mesure]
    # Recuperation des noms des colonens
    col_labels = df_nat.iloc[0]
    short_indic = dict_mesure_indic[mesure][0]
    date, valeur = get_kpi(departement, short_indic, mesure )
    context = { 'mesure' : mesure,
                'date': date,
                'indicateur': short_indic,
                'val': valeur,                
                'title_table_nat' : "Niveau National", 
                'title_table_reg' : "Niveau Régional", 
                'title_table_dep' : "Niveau Départemental",
                'lib_reg': reg_name, 
                'lib_dep': dep_name,
                'col_labels' : col_labels, 
                # Les 3 lignes suivantes permettent de générer des tabeaux avec uniquement les 3 derniers mois.
                # Prend en compte le cas ou il n'y a pas encore 3 mois de données
                'tbl_contents_nat': [{'cols' : list(df_nat.iloc[-i-1])} for i in range(min(len(df_nat)-1, 3))],
                'tbl_contents_reg': [{'cols' : list(df_reg.iloc[-i-1])} for i in range(min(len(df_reg)-1, 3))],
                'tbl_contents_dep': [{'cols' : list(df_dep.iloc[-(i+1)])} for i in range(min(len(df_dep)-1, 3))],
                }
    doc.render(context)
    name_file = "reports/test/content_page_{}.docx".format(mesure)
    doc.save(name_file)
    return name_file


def creation_fiche(dep):
    #departement: code departement 01:
    #On a les variables volet2mesures, all_charts
    reg = taxo_dep_df[taxo_dep_df['dep'] == dep].iloc[0]['reg']  # Code region
    if reg in L_reg_no_output:
        return False
    reg_name = taxo_reg_df[taxo_reg_df['reg'] == reg].iloc[0]['libelle']  # libelle
    dep_name = taxo_dep_df[taxo_dep_df['dep'] == dep].iloc[0]['libelle']
    name_fusion = creation_front_page(dep_name)
    for volet in list(volet2mesures.keys()):  # 3 itérations, dep_name, reg_name
        name_volet = creation_volet_page(volet)
        name_fusion = fusion2(name_fusion, name_volet, dep_name)
        liste_mesure = volet2mesures[volet]
        for mesure in liste_mesure:
            name_content = creation_content_page(all_charts_as_df, dep, reg, mesure, dep_name, reg_name)
            name_fusion = fusion2(name_fusion, name_content, dep_name)
 

def create_all_dep():
    list_all_dep = taxo_dep_df["dep"]
    for dep in list_all_dep:
        creation_fiche(dep)

t1 = time.clock()

create_all_dep()

t2 = time.clock()
print((t2-t1)/60)

