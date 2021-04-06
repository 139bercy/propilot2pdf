from docxtpl import DocxTemplate

doc = DocxTemplate("template/template_content_page.docx")
context = {'title_table_nat' : "Niveau National", 
            'title_table_dep' : "Niveau Departemental", 
            'title_table_reg' : "Niveau Regional", 
'col_labels' : ['Date','Nombre d"entreprises bénéficiaires', 'Nombre de TPE, PME, ETI bénéficiaires'],
'tbl_contents_dep': [
    {'cols': ['Novembre 2020', '4567', '836']},
    {'cols': ['Decembre 2020', '1306', '1306']},
    {'cols': ['Janvier 2021', '1647', '1645']},
    {'cols': ['Février 2021', '2057', '2054']},
    ],
'tbl_contents_reg': [
    {'cols': ['Novembre 2020', '836', '836']},
    {'cols': ['Decembre 2020', '1306', '1306']},
    {'cols': ['Janvier 2021', '1647', '9834']},
    {'cols': ['Février 2021', '2057', '2054']},
    ],
'tbl_contents_nat': [
    {'cols': ['Novembre 2020', '836', '836']},
    {'cols': ['Decembre 2020', '1306', '1306']},
    {'cols': ['Janvier 2021', '1647', '1645']},
    {'cols': ['Février 2021', '1', '2054']},
    ],
}
doc.render(context)
doc.save("generated_doc.docx")




from docx import Document

## Fusion de word
doc1 = Document('generated_doc.docx')
doc2 = Document('generated_doc.docx')
doc1.add_page_break()
doc1.add_page_break()
doc_new = Document()
for element in doc2.element.body:
    doc1.element.body.append(element)
doc1.save('new.docx')
